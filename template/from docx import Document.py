from docx import Document

# Create a new Word Document
doc = Document()
doc.add_heading('Software Development Life Cycle (SDLC) Tracker', 0)

# Add table with 7 columns and header row
table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'

# Define headers
headers = ["Phase", "Task Description", "Responsible Person", "Start Date", "End Date", "Status", "Feedback/Notes"]
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Define SDLC data
sdlc_data = [
    ["1. Requirement Analysis", "Gather business requirements, identify stakeholders.", "John Doe", "10-Jun-2025", "12-Jun-2025", "Completed", "Add more use cases for module X"],
    ["2. Planning", "Define scope, resources, and timeline.", "Jane Smith", "13-Jun-2025", "14-Jun-2025", "In Progress", "Timeline estimate seems optimistic"],
    ["3. Design", "Create architecture diagrams, database schema, UI mockups.", "Mark Lee", "15-Jun-2025", "17-Jun-2025", "Pending", "Waiting for planning finalization"],
    ["4. Development", "Write code, implement features.", "Dev Team", "18-Jun-2025", "25-Jun-2025", "Pending", ""],
    ["5. Testing", "Unit testing, integration testing, user acceptance testing.", "QA Team", "26-Jun-2025", "29-Jun-2025", "Pending", ""],
    ["6. Deployment", "Move code to production, configure infrastructure.", "Ops Team", "30-Jun-2025", "01-Jul-2025", "Pending", ""],
    ["7. Maintenance", "Monitor system, apply patches, handle bug reports.", "Support Team", "Ongoing", "Ongoing", "Not Started", ""]
]

# Add rows to the table
for row_data in sdlc_data:
    row_cells = table.add_row().cells
    for i, cell in enumerate(row_data):
        row_cells[i].text = cell

# Save the document
doc.save("SDLC_Tracker.docx")
